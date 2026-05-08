# -*- coding: utf-8 -*-
import os, sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR  = r"c:\Users\basem\OneDrive\سطح المكتب\Machine_project"
DOC_PATH = os.path.join(OUT_DIR, "Phase1_ML_Report_Updated.docx")

NAVY      = RGBColor(31, 56, 100)
RED_COLOR = RGBColor(192, 0, 0)
LBL = 'BDD7EE'
GRN = 'E2EFDA'
BP  = 10.5   # body point size

def shd(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    s = OxmlElement('w:shd')
    s.set(qn('w:val'), 'clear'); s.set(qn('w:color'), 'auto')
    s.set(qn('w:fill'), hex_color); tcPr.append(s)

def shade_row(row, c):
    for cell in row.cells: shd(cell, c)

def fmt(p, sb=0, sa=3, ls=1.0):
    pf = p.paragraph_format
    pf.space_before = Pt(sb); pf.space_after = Pt(sa); pf.line_spacing = ls

def h1(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True; r.font.size = Pt(13); r.font.color.rgb = NAVY
    fmt(p, sb=10, sa=3); return p

def h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True; r.font.size = Pt(11); r.font.color.rgb = NAVY
    fmt(p, sb=7, sa=2); return p

def h3(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True; r.font.size = Pt(BP)
    fmt(p, sb=4, sa=2); return p

def body(doc, text, bold=False, italic=False, centered=False, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.size = Pt(BP); r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = color
    if centered: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt(p, sb=0, sa=3); return p

def fmla(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.italic = True; r.font.size = Pt(BP)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER; fmt(p, sb=1, sa=1); return p

def mixed(doc, parts, centered=False):
    p = doc.add_paragraph()
    if centered: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for text, bold, italic, color in parts:
        r = p.add_run(text); r.font.size = Pt(BP); r.bold = bold; r.italic = italic
        if color: r.font.color.rgb = color
    fmt(p, sb=0, sa=3); return p

def tbl(doc, headers, rows, hi_rows=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for para in c.paragraphs:
            for r in para.runs: r.bold = True; r.font.size = Pt(9.5)
    shade_row(t.rows[0], LBL)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = str(val)
            for para in c.paragraphs:
                for r in para.runs: r.font.size = Pt(9.5)
        if hi_rows and ri in hi_rows: shade_row(t.rows[ri+1], GRN)
    p = doc.add_paragraph(); fmt(p, sb=0, sa=3); return t

def cm_tbl(doc, title, data):
    h3(doc, title)
    t = doc.add_table(rows=3, cols=3); t.style = 'Table Grid'
    t.rows[0].cells[0].text = ''; t.rows[0].cells[1].text = 'Pred Not 6'; t.rows[0].cells[2].text = 'Pred Is 6'
    t.rows[1].cells[0].text = 'Act. Not 6'; t.rows[2].cells[0].text = 'Act. Is 6'
    t.rows[1].cells[1].text = data[0][0]; shd(t.rows[1].cells[1], 'C6EFCE')
    t.rows[1].cells[2].text = data[0][1]; shd(t.rows[1].cells[2], 'FFC7CE')
    t.rows[2].cells[1].text = data[1][0]; shd(t.rows[2].cells[1], 'FFC7CE')
    t.rows[2].cells[2].text = data[1][1]; shd(t.rows[2].cells[2], 'C6EFCE')
    shade_row(t.rows[0], LBL)
    for row in t.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for r in para.runs: r.bold = True; r.font.size = Pt(9.5)
    p = doc.add_paragraph(); fmt(p, sb=0, sa=2)

# =============================================================================
doc = Document()
try: doc.styles['Normal'].font.name = 'Calibri'
except: pass
for sec in doc.sections:
    sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.5); sec.right_margin  = Cm(2.5)

# TITLE PAGE
p = doc.add_paragraph(); r = p.add_run('Ain Shams University')
r.bold = True; r.font.size = Pt(20); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; fmt(p, sb=60, sa=4)

p = doc.add_paragraph()
r = p.add_run('Faculty of Engineering — CAIE Program\nCSE382: Introduction to Machine Learning | Spring 2026')
r.font.size = Pt(12); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; fmt(p, sb=0, sa=20)

p = doc.add_paragraph()
r = p.add_run('Major Task Project — Phase 1 Technical Report')
r.bold = True; r.font.size = Pt(17); r.font.color.rgb = NAVY
p.alignment = WD_ALIGN_PARAGRAPH.CENTER; fmt(p, sb=0, sa=4)

p = doc.add_paragraph()
r = p.add_run('Binary Classification: MNIST Digit "6" Detection')
r.font.size = Pt(13); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; fmt(p, sb=0, sa=30)

p = doc.add_paragraph(); r = p.add_run('Team 1'); r.bold = True; r.font.size = Pt(13)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER; fmt(p, sb=0, sa=4)

team = [('23P0024','Adham Walid Said Zaki'),
        ('23P0248','Amir Tamer Mohamad Abdelrehim Mohamad'),
        ('23P0134','Mohamed Wael El Sayed Ali El Borai'),
        ('23P0246','Basem Walid Talaat Mansour Ahmed'),
        ('23P0049','Moaz Ahmed Mohamed Fathy Ahmed Ahmed Younis')]
tt = doc.add_table(rows=6, cols=2); tt.style='Table Grid'; tt.alignment=WD_TABLE_ALIGNMENT.CENTER
tt.rows[0].cells[0].text='Student ID'; tt.rows[0].cells[1].text='Full Name'
for c in tt.rows[0].cells:
    for r in c.paragraphs[0].runs: r.bold = True
shade_row(tt.rows[0], LBL)
for i,(sid,name) in enumerate(team):
    tt.rows[i+1].cells[0].text=sid; tt.rows[i+1].cells[1].text=name

p = doc.add_paragraph(); fmt(p, sb=20, sa=0)
p = doc.add_paragraph()
r = p.add_run('Due Date: Milestone 1 — Week 13')
r.font.size = Pt(11); r.font.color.rgb = RED_COLOR
p.alignment = WD_ALIGN_PARAGRAPH.CENTER; fmt(p, sb=0, sa=0)
doc.add_page_break()

# 1. PROBLEM DEFINITION
h1(doc, '1. Problem Definition')
body(doc,'The task is to build a binary image classifier on the MNIST handwritten digit dataset to detect instances of digit "6" against all others, formally defined as:')
fmla(doc,'y = 1  if digit is "6",    y = 0  otherwise')
body(doc,'The MNIST training set contains 60,000 labeled grayscale images (28x28 pixels). Input x is the flattened, normalised pixel vector; output y is the class label. Digit "6" represents approx. 10% of the dataset, creating a class-imbalanced problem with a Not-6 : 6 ratio of approximately 9:1.')

# 2. DATA PROCESSING
h1(doc, '2. Data Processing & Feature Engineering')
body(doc,'A consistent preprocessing pipeline is applied: (1) Flatten each 28x28 image to a 784-dim vector; (2) Normalise pixel intensities to [0, 1]; (3) Extract HOG features; (4) Apply PCA (50 components, fit on training data only); (5) Partition into train / validation / test sets. Class imbalance is addressed via sample weighting during training.')
h3(doc,'Feature Configurations')
mixed(doc,[('Raw Pixels — ',True,False,None),('Each image is flattened to a 784-dimensional vector (intensities in [0,1]). Preserves full spatial detail but is highly susceptible to the curse of dimensionality and pixel correlations.',False,False,None)])
mixed(doc,[('HOG (Histogram of Oriented Gradients) — ',True,False,None),('Gradient orientation histograms computed over 4x4-pixel cells (9 orientation bins), L2-Hys normalised over 2x2-cell blocks. Encodes the structural contours — closed loop and descending stroke — that distinguish "6", robust to illumination changes.',False,False,None)])
mixed(doc,[('PCA — ',True,False,None),('Principal Component Analysis retaining 50 components decorrelates features, discards low-variance noise, and makes Euclidean distance more meaningful. Fit on the training partition only to prevent data leakage.',False,False,None)])
mixed(doc,[('HOG + PCA (Primary Pipeline) — ',True,False,None),('HOG extracts meaningful shape features; PCA then compresses the HOG representation into 50 decorrelated components. This two-stage pipeline yields the best-performing input for all distance- and probability-based models.',False,False,None)])
body(doc,'Train / Validation / Test split: 48,000 / 12,000 / 10,000 samples (stratified).')

# 3. MODEL IMPLEMENTATIONS
h1(doc, '3. Model Implementations')

h2(doc, '3.1 Logistic Regression')
body(doc,'Models the posterior probability P(y=1|x) via a sigmoid-activated linear boundary:')
fmla(doc,'P(y = 1 | x) = sigma(w^T x + b) = 1 / (1 + exp(-(w^T x + b)))')
body(doc,'Optimised by minimising the Binary Cross-Entropy (Log Loss) using gradient descent:')
fmla(doc,'L = -(1/n) sum_i [ w_i * (y_i log(y_hat_i) + (1-y_i) log(1-y_hat_i)) ]')
body(doc,"Class imbalance is handled by scaling each sample's gradient update by its assigned class weight w_i, proportional to 1 / class frequency. Update rule: w <- w - eta * grad_L,  b <- b - eta * dL/db.")

h2(doc, '3.2 Gaussian Naive Bayes')
body(doc,"Applies Bayes' theorem under a conditional independence assumption. Each feature's class-conditional likelihood is modelled as a Gaussian:")
fmla(doc,'P(x_i | c) = (1 / sqrt(2*pi*sigma^2)) * exp( -(x_i - mu)^2 / (2*sigma^2) )')
fmla(doc,'log P(y=c | x) = log P(c) + sum_i log P(x_i | c)')
body(doc,'Log-probabilities are used throughout for numerical stability. A weight term log(w_c) is added to log-posteriors for class-imbalance correction. Parameters mu and sigma^2 are estimated from the training set (MLE).')

h2(doc, '3.3 Decision Tree')
body(doc,'Recursively partitions feature space by selecting the split (feature, threshold) that minimises weighted node impurity. Supported criteria:')
fmla(doc,'Gini(t) = 1 - sum_k p_k^2,    Entropy(t) = -sum_k p_k log(p_k)')
body(doc,'Class-weighted probability estimates: p_k = (w_k * n_k) / sum_j (w_j * n_j). Stopping conditions: max_depth, min_samples_split, node purity. Optimal hyperparameters found via grid search.')

h2(doc, '3.4 K-Nearest Neighbors (KNN)')
body(doc,"A non-parametric, lazy learning algorithm that stores the entire training set and classifies each test sample by majority vote among its k closest neighbours in feature space. The Euclidean distance between two feature vectors x and x' is:")
fmla(doc,"d(x, x') = sqrt( sum_i (x_i - x'_i)^2 )")
body(doc,'Distance-weighted voting is used throughout, assigning each neighbour a vote weight of 1/(d_i + epsilon) where epsilon = 10^-5 prevents division by zero for exact matches. This scheme suppresses the influence of distant neighbours and consistently outperforms uniform voting on imbalanced datasets. Odd values of k are used to eliminate ties in binary classification.')

h2(doc, '3.5 Linear SVM (from scratch)')
body(doc,'Implemented entirely from scratch using mini-batch sub-gradient descent on the soft-margin primal objective. Labels encoded as y in {-1, +1}.')
h3(doc,'Objective Function (Class-Weighted Soft-Margin Hinge Loss)')
fmla(doc,'J(w, b) = (1/2)||w||^2 + C * sum_i s_i * max(0, 1 - y_i(w^T x_i + b))')
body(doc,'The regularisation term (1/2)||w||^2 maximises the geometric margin 2/||w||; the hinge loss penalises margin violations; C controls the bias-variance tradeoff. Per-sample weight s_i corrects for class imbalance:')
fmla(doc,'s_i = n / (2 * n_class(i))')
h3(doc,'Sub-Gradient Derivation')
body(doc,'For violating samples (y_i(w^T x_i + b) < 1):')
fmla(doc,'dJ/dw = w - C * sum_{i in violators} s_i * y_i * x_i')
fmla(doc,'dJ/db = -C * sum_{i in violators} s_i * y_i')
body(doc,'For non-violating points, the hinge gradient is zero and only the regulariser contributes (dJ/dw = w).')
h3(doc,'Parameter Update (Mini-Batch Sub-Gradient Descent)')
fmla(doc,'w <- w - eta * dJ/dw,    b <- b - eta * dJ/db')
body(doc,'Batch size = 64; data shuffled at the start of each epoch; w, b initialised to zero. Prediction: f(x) = w^T x + b >= 0 -> class +1 ("is a 6"), otherwise -1.')

# 4. EXPERIMENTAL RESULTS
h1(doc, '4. Experimental Results')

# 4.1 LR
h2(doc, '4.1 Logistic Regression')
h3(doc,'Hyperparameter Tuning — Best Weight Configuration')
tbl(doc,['Feature Config','Best Weight','Precision (6)','Recall (6)','F1 (6)','Accuracy'],
    [['Raw Pixels','Custom (1:5)','0.786','0.943','0.857','0.970'],
     ['PCA only','No Weights','0.894','0.849','0.871','0.976'],
     ['HOG only (best)','Custom (1:5)','0.943','0.966','0.954','0.991'],
     ['HOG + PCA','No Weights','0.988','0.880','0.931','0.988']],hi_rows=[2])
h3(doc,'5-Fold Cross-Validation (HOG only, weight = 1:5)')
body(doc,'Average F1 = 0.9444. Consistent scores across all folds confirm stable generalisation.')
h3(doc,'Best Configuration Analysis')
body(doc,'The optimal pipeline is HOG Only with a 1:5 custom weight (Accuracy 0.991, F1 0.954). HOG features contain localised, spatially meaningful shape information that a linear boundary can exploit cleanly. In uncompressed HOG space, a 1:5 class weight correctly rebalances minority-class sensitivity without overcorrection. When PCA is applied, features collapse into dense, blended components where adding class weights overcorrects the boundary, destroying precision — this is why HOG+PCA without weights achieves lower recall (0.880). The confusion matrix for the best config shows 925/958 true positives and only 56 false positives across 9,042 non-target digits.')

# 4.2 GNB
h2(doc, '4.2 Gaussian Naive Bayes')
h3(doc,'Weight Tuning (Validation F1, HOG+PCA)')
tbl(doc,['Class Weight','1','2','5','7','10 (best)'],[['Validation F1','0.9309','0.9405','0.9510','0.9525','0.9536']])
body(doc,'5-fold cross-validation average F1 = 0.9530, confirming stable generalisation.')
h3(doc,'Final Test Results — HOG + PCA, weight = 10')
tbl(doc,['Class','Precision','Recall','F1-score'],[['Not 6','0.99','1.00','1.00'],['Is 6','0.98','0.94','0.96']])
body(doc,'Overall Accuracy: 0.9918')
h3(doc,'Feature Configuration Comparison')
tbl(doc,['Feature Config','Precision (6)','Recall (6)','F1 (6)','Accuracy'],
    [['Raw Pixels','0.24','0.99','0.39','0.699'],['PCA only','0.85','0.95','0.90','0.979'],
     ['HOG only','0.44','0.98','0.61','0.879'],['HOG + PCA (best)','0.98','0.94','0.96','0.992']],hi_rows=[3])
body(doc,"GNB's performance is almost entirely determined by how well the input satisfies the conditional independence assumption. Raw pixels violate it catastrophically (F1 = 0.39 due to strong pixel correlations). PCA alone decorrelates features, jumping F1 to 0.90. HOG alone improves shape encoding but leaves correlated cells in place (F1 = 0.61). HOG+PCA satisfies both requirements simultaneously — near-zero off-diagonal correlations and broadly Gaussian feature distributions — yielding the best F1 of 0.96. Weight w=10 is necessary because GNB's generative prior is more strongly affected by class imbalance than discriminative models.")

# 4.3 DT
h2(doc, '4.3 Decision Tree')
h3(doc,'Key Hyperparameters')
tbl(doc,['Parameter','Description'],
    [['max_depth','Controls bias-variance tradeoff; low values underfit, high values overfit.'],
     ['min_samples_split / leaf','Regularise against overfitting on rare patterns; higher = smoother boundary.'],
     ['criterion (Gini / Entropy)','Entropy tends to produce more balanced splits on imbalanced datasets.']])
h3(doc,'Class Weight Comparison & Cross-Validation (HOG features)')
tbl(doc,['Configuration','Precision (6)','Recall (6)','F1 (6)','Accuracy'],
    [['No weights','0.94','0.87','0.91','0.98'],['Weight = 5 (best)','0.90','0.91','0.91','0.98']],hi_rows=[1])
body(doc,'5-fold cross-validation (HOG, weight = 5) average F1 = 0.9151. Folds ranged 0.908 – 0.931, confirming consistent generalisation.')
h3(doc,'Final Test Results — HOG only, weight = 5')
tbl(doc,['Class','Precision','Recall','F1-score'],[['Not 6','0.99','0.99','0.99'],['Is 6','0.90','0.91','0.91']])
body(doc,'Overall Accuracy: 0.98')
h3(doc,'Feature Configuration Comparison')
tbl(doc,['Feature Config','Precision (6)','Recall (6)','F1 (6)','Accuracy'],
    [['Raw Pixels','0.92','0.86','0.89','0.98'],['PCA only','0.83','0.86','0.84','0.97'],
     ['HOG only (best)','0.90','0.91','0.91','0.98'],['HOG + PCA','0.92','0.88','0.90','0.98']],hi_rows=[2])
body(doc,'Decision trees split on one feature at a time; each feature must carry localised, interpretable meaning to produce clean thresholds. HOG features are ideal: each value encodes the dominant edge orientation in a specific spatial cell, and digit "6"\'s characteristic loop and stem activate predictable cells. PCA destroys this locality by mixing all image regions into each component, making it the worst configuration (F1 = 0.84). HOG alone is uniquely the best for Decision Trees — the only model in this study where PCA hurts rather than helps.')

# 4.4 KNN -- REAL NUMBERS
h2(doc, '4.4 K-Nearest Neighbors (KNN)')
h3(doc,'Hyperparameter Tuning (HOG+PCA features, distance weighting)')
body(doc,'The training set was split 80/20 (stratified) into 48,000-sample training and 12,000-sample validation partitions. Odd k values evaluated to eliminate ties:')
tbl(doc,['k','Validation Error','Accuracy'],
    [['1','0.0032','0.9968'],['3','0.0032','0.9968'],['5','0.0032','0.9968'],
     ['7 (optimal)','0.0030','0.9970'],['9','0.0032','0.9968'],
     ['11','0.0033','0.9967'],['13','0.0034','0.9966'],
     ['15','0.0033','0.9967'],['17','0.0033','0.9967']],hi_rows=[3])
body(doc,'The error curve forms a clear elbow at k = 7, which was selected as the optimal value.')
h3(doc,'5-Fold Cross-Validation (k = 7, distance weighting)')
body(doc,'Performed on a 10,000-sample subset of the training partition (full 48k cross-validation requires approx. 30 min/fold):')
tbl(doc,['Fold','F1 Score (Class "6")'],
    [['1','0.9846'],['2','0.9673'],['3','0.9849'],['4','0.9596'],['5','0.9793'],['Average','0.9751']],hi_rows=[5])
body(doc,'Average F1 of 0.9751 confirms consistent generalisation; fold variance (0.9596–0.9849) reflects the limited 10k training subset per fold.')
h3(doc,'Final Test Results — HOG+PCA, k = 7, distance weighting')
tbl(doc,['','Not 6','Is 6','Macro Avg'],
    [['Precision','1.00','0.98','0.99'],['Recall','1.00','0.99','0.99'],
     ['F1-Score','1.00','0.98','0.99'],['Accuracy','—','—','0.9969']])
h3(doc,'Feature Configuration Comparison')
tbl(doc,['Feature Config','Precision (6)','Recall (6)','F1 (6)','Accuracy'],
    [['Raw Pixels†','0.9629','0.9760','0.9694','0.9941'],
     ['PCA only','0.9813','0.9885','0.9849','0.9971'],
     ['HOG only†','0.9711','0.9812','0.9761','0.9954'],
     ['HOG + PCA (best)','0.9803','0.9875','0.9839','0.9969']],hi_rows=[3])
body(doc,'†10,000-sample training subset used for computational tractability.')
h3(doc,'Analysis')
body(doc,'Raw pixels suffer from the curse of dimensionality: in 784-dimensional space, Euclidean distances become nearly uniform across samples, destroying the neighbourhood signal. PCA alone alleviates this by reducing to 50 discriminative components, but captures global intensity variance rather than structural shape. HOG alone greatly improves clustering — two handwritten "6"s with different ink styles share similar HOG vectors because their edge structures match — but minor cell redundancy persists without PCA. HOG+PCA is the best configuration: HOG builds tight, shape-consistent clusters while PCA removes remaining redundancy and keeps only the 50 most informative directions. Distance-weighted voting further benefits this imbalanced task: a very close "6" neighbour outweighs several distant non-"6" neighbours, reducing false positives without any explicit reweighting parameter to tune.')

# 4.5 SVM
h2(doc, '4.5 Linear SVM (from scratch)')
h3(doc,'Hyperparameter Tuning — Grid Search (C and learning rate)')
tbl(doc,['C','Learning Rate','Validation F1'],
    [['0.01','0.001','0.8354'],['0.1','0.001','0.8685'],
     ['1.0','0.0005 (best)','0.8704'],['10.0','0.0005','0.8379']],hi_rows=[2])
h3(doc,'Final Test Results — HOG only, C = 1.0, lr = 0.0005')
tbl(doc,['Class','Precision','Recall','F1-score'],[['Not 6','1.000','0.990','0.995'],['Is 6','0.951','0.987','0.969']])
body(doc,'Overall Accuracy: 0.9939')
h3(doc,'Feature Configuration Comparison')
tbl(doc,['Feature Config','Precision (6)','Recall (6)','F1 (6)','Accuracy'],
    [['Raw Pixels','0.79','0.97','0.87','0.97'],['PCA only','0.77','0.97','0.86','0.97'],
     ['HOG only (best)','0.95','0.99','0.97','0.9939'],['HOG + PCA','0.93','0.99','0.95','0.99']],hi_rows=[2])
body(doc,"The SVM's class-weighted hinge loss drives consistently high recall (0.97–0.99) across all feature configurations. HOG is the best configuration (F1 = 0.969): HOG's edge-orientation encoding gives the SVM a structured, locally-meaningful feature space where the margin-maximising objective can draw a tight boundary around digit \"6\"'s distinctive loop and stroke. Raw Pixels perform surprisingly poorly (F1 = 0.87) compared to the SVM's other configs — in 784-dimensional raw space, class-weighted hinge loss still finds a discriminative hyperplane, but correlated pixel features force the margin to be noisier and the boundary less precise. PCA alone is nearly as weak (F1 = 0.86): global pixel projections lose the structural signal entirely. HOG+PCA (F1 = 0.95) is competitive but slightly below HOG alone — PCA compression discards a small amount of edge-orientation detail that the SVM's margin objective would otherwise exploit.")

# 4.6 CONFUSION MATRICES
h2(doc, '4.6 Confusion Matrices (Best Configuration per Model)')
body(doc,'Green = correct predictions (TN / TP), Red = errors (FP / FN).')
cm_tbl(doc,'Logistic Regression',[['8986','56'],['33','925']])
cm_tbl(doc,'Gaussian Naive Bayes',[['9026','16'],['62','896']])
cm_tbl(doc,'Decision Tree',[['8950','92'],['88','870']])
cm_tbl(doc,'SVM (HOG only)',[['8993','49'],['12','946']])
cm_tbl(doc,'KNN (HOG+PCA, k=7)',[['9023','19'],['11','947']])
body(doc,'LR (HOG only, 1:5 weights): 925 TP, 56 FP, 33 FN — a well-balanced boundary. GNB (HOG+PCA, w=10): only 16 FP (highest precision of all models), but 62 FN reflects the generative model\'s weaker recall. DT (HOG only, w=5): weakest profile with 92 FP and 88 FN. SVM (HOG only, C=1.0): standout — only 12 FN (fewest of all models), confirming that HOG features combined with class-weighted hinge loss maximally preserve minority-class recall. KNN (HOG+PCA, k=7): 947 TP, 19 FP, 11 FN — highly precise boundary enabled by distance-weighted voting in a well-structured feature space.')

# 5. MODEL COMPARISON
h1(doc, '5. Model Comparison and Discussion')
body(doc,'All metrics derived directly from test-set confusion matrices or final evaluation runs. Models ranked by F1 for class "6".')
tbl(doc,['Model','Precision (6)','Recall (6)','F1 (6)','Accuracy','Best Features'],
    [['KNN (k=7, dist.)','0.980','0.988','0.984','0.9969','HOG+PCA'],
     ['Linear SVM','0.95','0.99','0.97','0.9939','HOG only'],
     ['Gaussian Naive Bayes','0.98','0.94','0.96','0.9918','HOG+PCA'],
     ['Logistic Regression','0.943','0.966','0.954','0.991','HOG only'],
     ['Decision Tree (w=5)','0.90','0.91','0.91','0.98','HOG only']],hi_rows=[0])
body(doc,'KNN (F1 = 0.984, Acc = 0.9969) is the best overall model. HOG+PCA creates tight, well-separated clusters where distance-weighted voting naturally handles the 9:1 class imbalance without any explicit reweighting parameter.')
body(doc,"The SVM ranks 2nd (F1 = 0.97, Acc = 0.9939) using HOG features — and unlike KNN, does not need PCA. HOG's edge-orientation features give the hinge-loss objective a clean feature space, producing only 12 false negatives, the lowest of any model. This best-in-class recall (0.99) makes the SVM the preferred choice for any recall-critical application.")
body(doc,'GNB ranks 3rd (F1 = 0.96, Acc = 0.9918) with the highest precision (0.98) among all models, producing only 16 FP while missing 62 positive samples. Its near-decorrelated, near-Gaussian HOG+PCA features maximally satisfy the conditional independence assumption.')
body(doc,"Logistic Regression (F1 = 0.954, Acc = 0.991) is the only model where HOG+PCA hurts: PCA compresses features into dense blended components where the 1:5 class weight overcorrects the boundary, collapsing precision. HOG alone with 1:5 weighting gives the cleanest linear separation.")
body(doc,'Decision Tree is the weakest performer (F1 = 0.91, Acc = 0.98) with the highest combined error count (88 FN + 92 FP = 180 total errors, versus the SVM\'s 61). Its axis-aligned splits cannot capture digit "6"\'s curved manifold, and it is the only model where PCA degrades performance by destroying the spatial locality each threshold split requires.')

# 6. CONCLUSION
h1(doc, '6. Conclusion')
body(doc,'This report presented a complete binary classification pipeline for MNIST digit "6" detection, implementing five algorithms from scratch — Logistic Regression, Gaussian Naive Bayes, Decision Tree, KNN, and Linear SVM — with HOG+PCA as the primary feature engineering pipeline.')
body(doc,'KNN (k=7, distance weighting, HOG+PCA) achieves the best overall balance with F1 = 0.984 and accuracy = 0.9969. The Linear SVM is the standout for recall-critical applications (only 12 FN, F1 = 0.97). HOG+PCA is the single most impactful preprocessing step; the Decision Tree is the only exception, benefiting from HOG alone due to its reliance on locally-interpretable, spatially-meaningful feature splits.')
body(doc,'Phase 2 will scale this pipeline to 10-class classification and incorporate pretrained CNN feature extraction, ensemble methods, and regularisation analysis.')

doc.save(DOC_PATH)
print("Saved ->", DOC_PATH)

# Page count
try:
    import win32com.client
    w = win32com.client.Dispatch('Word.Application')
    w.Visible = False
    d = w.Documents.Open(DOC_PATH)
    pages = d.ComputeStatistics(2)
    d.Close(False); w.Quit()
    print(f"Page count: {pages}")
except Exception as e:
    print(f"(Could not count pages: {e})")
